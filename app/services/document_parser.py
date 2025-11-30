import os
import tempfile
import asyncio
from typing import Dict, Tuple, Optional
import aiohttp
import aiofiles
from loguru import logger

# Document parsing libraries
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import io


class DocumentParser:
    """
    Service for parsing various document formats (PDF, DOCX, Images)
    """

    def __init__(self):
        self.supported_formats = ["pdf", "docx", "doc", "jpg", "jpeg", "png", "tiff"]

    async def parse_from_url(
        self, url: str, file_type: str, language: str = "en"
    ) -> Tuple[str, float]:
        """
        Download and parse document from URL
        """
        start_time = asyncio.get_event_loop().time()

        try:
            # Download file
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    if response.status != 200:
                        raise ValueError(
                            f"Failed to download file from URL: {response.status}"
                        )

                    file_content = await response.read()

            # Create temporary file
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=f".{file_type}"
            ) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name

            try:
                # Parse the file
                extracted_text = await self._parse_file(temp_path, file_type, language)
                processing_time = asyncio.get_event_loop().time() - start_time

                return extracted_text, processing_time

            finally:
                # Clean up temporary file
                os.unlink(temp_path)

        except Exception as e:
            logger.error(f"Failed to parse document from URL {url}: {e}")
            raise

    async def parse_local_file(
        self, file_path: str, file_type: str
    ) -> Tuple[str, float]:
        """
        Parse local document file
        """
        start_time = asyncio.get_event_loop().time()

        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            extracted_text = await self._parse_file(file_path, file_type)
            processing_time = asyncio.get_event_loop().time() - start_time

            return extracted_text, processing_time

        except Exception as e:
            logger.error(f"Failed to parse local file {file_path}: {e}")
            raise

    async def _parse_file(
        self, file_path: str, file_type: str, language: str = "en"
    ) -> str:
        """
        Internal method to parse different file formats
        """
        try:
            if file_type.lower() in ["pdf"]:
                return await self._parse_pdf(file_path)
            elif file_type.lower() in ["docx", "doc"]:
                return await self._parse_docx(file_path)
            elif file_type.lower() in ["jpg", "jpeg", "png", "tiff"]:
                return await self._parse_image(file_path, language)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error parsing file {file_path} as {file_type}: {e}")
            raise

    async def _parse_pdf(self, file_path: str) -> str:
        """
        Parse PDF file using pdfplumber
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if not text.strip():
                logger.warning(
                    f"PDF {file_path} returned no text - may be scanned/image-based"
                )
                # Fallback: try OCR if no text found
                return await self._parse_image(file_path, "en")

            return text.strip()

        except Exception as e:
            logger.error(f"PDF parsing failed for {file_path}: {e}")
            raise

    async def _parse_docx(self, file_path: str) -> str:
        """
        Parse DOCX file using python-docx
        """
        try:
            doc = Document(file_path)
            text = "\n".join(
                [
                    paragraph.text
                    for paragraph in doc.paragraphs
                    if paragraph.text.strip()
                ]
            )

            if not text.strip():
                # Try to extract from tables as well
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"DOCX parsing failed for {file_path}: {e}")
            raise

    async def _parse_image(self, file_path: str, language: str = "en") -> str:
        """
        Parse image file using OCR (Tesseract)
        """
        try:
            # Configure Tesseract for better results
            custom_config = r"--oem 3 --psm 6"

            if language == "ar":
                custom_config += " -l ara"
            else:
                custom_config += " -l eng"

            # Open and process image
            image = Image.open(file_path)

            # Preprocess image for better OCR
            image = self._preprocess_image(image)

            # Perform OCR
            text = pytesseract.image_to_string(image, config=custom_config)

            return text.strip()

        except Exception as e:
            logger.error(f"Image OCR failed for {file_path}: {e}")
            raise

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image to improve OCR accuracy
        """
        try:
            # Convert to grayscale if needed
            if image.mode != "L":
                image = image.convert("L")

            # You can add more preprocessing here:
            # - Noise reduction
            # - Contrast enhancement
            # - Deskewing
            # - etc.

            return image

        except Exception as e:
            logger.warning(f"Image preprocessing failed, using original: {e}")
            return image

    def validate_file_type(self, file_type: str) -> bool:
        """
        Validate if file type is supported
        """
        return file_type.lower() in self.supported_formats

    async def get_file_metadata(self, file_path: str) -> Dict:
        """
        Extract basic file metadata
        """
        try:
            stat = os.stat(file_path)
            return {
                "file_size": stat.st_size,
                "created_time": stat.st_ctime,
                "modified_time": stat.st_mtime,
            }
        except Exception as e:
            logger.error(f"Failed to get file metadata for {file_path}: {e}")
            return {}


# Singleton instance
document_parser = DocumentParser()
